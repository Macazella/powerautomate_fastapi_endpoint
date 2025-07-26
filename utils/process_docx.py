import os
from docx import Document
from fastapi import UploadFile

async def process_docx_file(file: UploadFile) -> dict:
    contents = await file.read()
    with open("temp.docx", "wb") as f:
        f.write(contents)
    
    doc = Document("temp.docx")
    texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    os.remove("temp.docx")
    return {
        "tipo": "docx",
        "parrafos_detectados": len(doc.paragraphs),
        "contenido": texto[:500] + "..." if len(texto) > 500 else texto
    }
